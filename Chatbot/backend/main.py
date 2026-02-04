# main.py
from fastapi.responses import StreamingResponse
from fastapi import FastAPI, HTTPException, Query, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from typing import List
import os
import base64
from pathlib import Path
from docx import Document
from pydantic import BaseModel
from groq import Groq
import uuid
from datetime import datetime
from dotenv import load_dotenv
import json
import re
import logging
import fitz  # PyMuPDF
from PIL import Image as PILImage
import requests
import io

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

app = FastAPI()
GROQ_MODEL = "llama-3.3-70b-versatile"
GROQ_VISION_MODEL = "llama-3.2-11b-vision-preview"

# CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"
CHAT_HISTORY_DIR = BASE_DIR / "chat_history"

# Create chat history directory if it doesn't exist
CHAT_HISTORY_DIR.mkdir(exist_ok=True)

@app.get("/api/templates/{category}")
def list_templates(category: str):
    category_path = TEMPLATES_DIR / category
    if not category_path.exists() or not category_path.is_dir():
        raise HTTPException(status_code=404, detail="Category not found")
    files = [f.name for f in category_path.glob("*.docx")]
    return {"templates": files}

@app.get("/api/template")
def get_template(category: str = Query(...), name: str = Query(...)):
    file_path = TEMPLATES_DIR / category / name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    with open(file_path, "rb") as f:
        encoded = base64.b64encode(f.read()).decode("utf-8")
    return {"base64": encoded}

@app.get("/api/template/sections")
def get_template_sections(category: str = Query(...), name: str = Query(...)):
    file_path = TEMPLATES_DIR / category / name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    doc = Document(file_path)
    sections = [para.text.strip() for para in doc.paragraphs if para.text.strip().lower().startswith("template for")]
    if not sections:
        sections = ["Full Document"]
    return {"sections": sections}

load_dotenv()
groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

class AIStartRequest(BaseModel):
    category: str
    subtype: str | None = None
    user_input: str

class ChatRequest(BaseModel):
    message: str
    history: List[dict] = []
    session_id: str | None = None

def load_metadata(category: str, subtype: str | None = None):
    if subtype:
        target_path = TEMPLATES_DIR / category / subtype / "metadata.json"
    else:
        target_path = TEMPLATES_DIR / category / "metadata.json"
    
    if not target_path.exists():
        logging.error(f"Metadata file not found at expected path: {target_path.absolute()}")
        raise HTTPException(
            status_code=404, 
            detail=f"metadata.json not found for category='{category}' subtype='{subtype}'. Expected at: {target_path}"
        )
    
    try:
        with open(target_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError:
        logging.error(f"Failed to decode JSON from {target_path}")
        raise HTTPException(status_code=500, detail="metadata.json is not a valid JSON file")

@app.post("/api/ai/start")
def start_ai_flow(data: AIStartRequest):
    try:
        templates = load_metadata(data.category, data.subtype)
        
        if not templates:
            raise HTTPException(status_code=404, detail="No templates found in metadata.")
        for t in templates:
            if not all(k in t for k in ("title", "summary", "filename")):
                raise HTTPException(status_code=500, detail="Invalid metadata format")

        summaries_str = "\n\n".join([
            f"Title: {t['title']}\nSummary: {t['summary']}\nFilename: {t['filename']}"
            for t in templates
        ])

        selection_prompt = (
        "You are an expert legal assistant helping users find the most suitable document template.\n"
        f"A user described their issue as:\n\n{data.user_input.strip()}\n\n"
        "Here are the available legal document templates:\n\n"
        f"{summaries_str}\n\n"
        "Your task:\n"
        "— Select the single most suitable template.\n"
        "— Respond ONLY with the value of the 'filename' field (e.g. template_for_unreasonableness).\n"
        "— Do NOT explain your choice.\n"
        "— If multiple templates could apply, choose the best one based on the user's description.\n"
        "— If unsure, pick the closest reasonable match anyway — do NOT say 'none match'."
    )

        try:
            response = groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {"role": "system", "content": "You are a helpful legal assistant..."},
                    {"role": "user", "content": selection_prompt},
                ],
                temperature=0.3,
                max_tokens=60,
            )
        
        except Exception as e:
            logging.error(f"Groq API error during template selection: {str(e)}")
            # Fallback: Pick the first template if AI is down
            logging.info("Falling back to first available template.")
            return {
                "question": "The AI service is temporarily unavailable. Let's start with the basics: What is your name and address?",
                "filename": f"{data.subtype}/{templates[0]['filename']}" if data.subtype else templates[0]['filename']
            }

        selected_filename = response.choices[0].message.content.strip()
        logging.info(f"Selected Filename from GPT: {selected_filename}")

        matched_template = next(
            (t for t in templates if t["filename"].strip().lower() == selected_filename.lower()),
            None
        )

        if not matched_template:
            raise HTTPException(status_code=404, detail="Template match not found in metadata")

        doc_filename = f"{selected_filename}.docx"
        logging.info(f"Selected Template File: {doc_filename}")

        if data.subtype:
            doc_path = TEMPLATES_DIR / data.category / data.subtype / doc_filename
        else:
            doc_path = TEMPLATES_DIR / data.category / doc_filename

        if not doc_path.exists():
            raise HTTPException(status_code=404, detail="Selected template file not found")

        doc = Document(doc_path)
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        question_prompt = (
            "You are a professional legal assistant helping a user complete a legal document.\n"
            "You are given the full text of the template. Read it carefully and identify all placeholders or gaps that must be completed by the user (e.g. [insert full address], empty lines, bullet point options, or areas left blank for details).\n\n"
            "Ask questions one at a time to gather the exact information needed to fill in these blanks. Start with the most essential or obvious missing fields.\n"
            "Make each question clear, simple, and specific — just like you're guiding someone through a form.\n"
            "If there are multiple options in a section (e.g. a, b, c), ask follow-up questions to help the user choose the correct one.\n"
            "Do not explain the document. Just act like a legal assistant who knows what details are needed and asks for them naturally, one by one.\n"
            "Avoid asking for contact info unless the template explicitly requires it.\n\n"
            "Once the necessary information has been collected, the document will be auto-completed and downloaded by the user."
        )

        try:
            q_response = groq_client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[
                    {"role": "system", "content": question_prompt},
                    {"role": "user", "content": full_text},
                ],
                temperature=0.3,
                max_tokens=150,
            )
        except Exception as e:
            logging.error(f"Groq API error during first question generation: {str(e)}")
            return {
                "question": "I've selected a template for you. To begin, could you provide your full legal name and current address?",
                "filename": f"{data.subtype}/{selected_filename}" if data.subtype else selected_filename
            }

        question = q_response.choices[0].message.content.strip()
        return {
        "question": question,
        "filename": f"{data.subtype}/{selected_filename}" if data.subtype else selected_filename
    }

    except Exception as e:
        logging.exception("Unexpected server error")
        raise HTTPException(status_code=500, detail=str(e))
    
class AINextRequest(BaseModel):
    category: str
    filename: str
    messages: List[dict]

@app.post("/api/ai/next")
def ai_next_question(data: AINextRequest):
    try:
        subtype_path = data.filename.split("/", 1)
        if len(subtype_path) == 2:
            subfolder, filename = subtype_path
            file_path = TEMPLATES_DIR / data.category / subfolder / f"{filename}.docx"
        else:
            file_path = TEMPLATES_DIR / data.category / f"{data.filename}.docx"

        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Template not found")

        doc = Document(file_path)
        template_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        system_prompt = (
            "You are a professional legal assistant continuing a session to help a user complete a legal document.\n"
            "You have access to the full document template and the conversation history.\n"
            "Identify any remaining placeholders (like [insert...], blank lines, bullet point choices, or missing details).\n"
            "Ask ONE specific, clear question at a time to gather that missing information.\n"
            "If all placeholders are filled, respond with __COMPLETE__ to signal the document is ready for generation.\n"
            "Do not explain or summarize the document — focus only on gathering the required inputs naturally and efficiently."
        )

        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"This is the template:\n\n{template_text}"},
                *data.messages
            ],
            temperature=0.3,
            max_tokens=300
        )

        reply = response.choices[0].message.content.strip()
        return {"reply": reply}

    except Exception as e:
        logging.error(f"Groq API error during Q&A: {str(e)}")
        return {"reply": "I apologize, but I'm having trouble connecting to my brain right now. Please try again in a moment, or ensure all required fields are provided."}

    except Exception as e:
        logging.exception("Unexpected server error during /api/ai/next")
        raise HTTPException(status_code=500, detail=str(e))

class AICompleteRequest(BaseModel):
    category: str
    filename: str
    messages: List[dict]

def extract_answers(messages: List[dict]):
    return [messages[i]["content"] for i in range(1, len(messages), 2) if messages[i]["role"] == "user"]

@app.get("/api/categories")
def list_categories():
    categories = {}
    for cat in TEMPLATES_DIR.iterdir():
        if not cat.is_dir(): continue
        subtypes = [sub.name for sub in cat.iterdir() if sub.is_dir() and ((sub / "metadata.json").exists() or any(f.suffix == ".docx" for f in sub.glob("*.docx")))]
        if subtypes or any(f.suffix == ".docx" for f in cat.glob("*.docx")) or (cat / "metadata.json").exists():
            categories[cat.name] = subtypes
    return categories

@app.post("/api/ai/complete")
def complete_template(data: AICompleteRequest):
    filename = data.filename.replace(".docx", "")
    file_path = TEMPLATES_DIR / data.category / f"{filename}.docx"
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")
    doc = Document(file_path)

    # GPT-driven fill‑in logic
    raw_template = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    chat_log = "\n".join(f"{m['role']}: {m['content']}" for m in data.messages)

    fill_prompt = f"""
You are a professional legal assistant.
Fill every gap in the document below (brackets, underlines, dotted blanks) using only the conversation data.

DOCUMENT TEMPLATE:
{raw_template}

CONVERSATION:
{chat_log}

Return JSON with:
  nextQuestion: <string or null>
  filledDocument: <complete text or null>
"""

    try:
        resp = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role":"system","content":"You fill in and ask for missing info. Your output must be valid JSON."},
                {"role":"user","content": fill_prompt}
            ],
            temperature=0.2,
            max_tokens=2000
        )
        content = resp.choices[0].message.content.strip()
        
        # Groq sometimes wraps JSON in markdown blocks
        if "```json" in content:
            content = content.split("```json")[-1].split("```")[0].strip()
        
        result = json.loads(content)

        if result.get("nextQuestion"):
            return {"nextQuestion": result["nextQuestion"]}
    except Exception as e:
        logging.error(f"Error during completion: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to generate document. Please ensure all details are provided.")

    # build final .docx from result["filledDocument"]
    from io import BytesIO
    out_doc = Document()
    for line in result["filledDocument"].split("\n"):
        out_doc.add_paragraph(line)
    buf = BytesIO()
    out_doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition":f"attachment; filename=filled_{filename}.docx"}
    )

@app.post("/api/chat")
def legal_chat(data: ChatRequest):
    """
    Guardian - Legal Information Assistant for Indian Law.
    Provides clear, accurate, and responsible legal information.
    Automatically saves chat history.
    """
    try:
        # Generate session_id if not provided
        session_id = data.session_id or str(uuid.uuid4())
        
        system_prompt = (
            "You are Guardian, a Legal Information Assistant for Indian law.\n"
            "Your role is to provide clear, simple, accurate, and responsible legal information to users. "
            "You are not a lawyer and you do not give professional legal advice.\n\n"
            "🎯 CORE OBJECTIVES:\n"
            "- Answer legal questions based on Indian law.\n"
            "- Use simple language suitable for non-lawyers.\n"
            "- Be factually accurate, updated, and neutral.\n"
            "- Never encourage illegal, harmful, or unethical behavior.\n\n"
            "📚 LEGAL SCOPE & ACCURACY RULES:\n"
            "- Criminal law: Use IPC (Indian Penal Code) for offences committed before 1 July 2024. Use BNS (Bharatiya Nyaya Sanhita) for offences committed on or after 1 July 2024.\n"
            "- Explain Civil, family, property, labour, and cyber laws at a high level.\n"
            "- If facts are insufficient, explain possible sections and state that the final decision depends on the court.\n"
            "- Never say a valid Indian law 'does not exist' if it is legally recognized (e.g., BNS).\n\n"
            "🚫 STRICT SAFETY RULES:\n"
            "- Do NOT provide: Instructions to commit crimes, advice to escape punishment, or guidance for violence/fraud.\n"
            "- If a user asks how to avoid punishment or expresses violent intent: Refuse to assist with wrongdoing, provide high-level legal consequences ONLY, and encourage lawful behavior.\n\n"
            "🗣️ RESPONSE STYLE:\n"
            "- Be concise but complete. Use bullet points or steps where helpful.\n"
            "- Avoid unnecessary legal jargon. Be respectful and neutral.\n"
            "- Do NOT sound robotic or threatening.\n\n"
            "- 🔒 DOMAIN RESTRICTION RULE:\n"
            "- Answer ONLY questions related to Indian law and legal matters."
            "- If a question is non-legal, general, personal, technical, or unrelated to law, clearly state that you cannot answer it and ask the user to reframe the query as a legal question."
            "- Do NOT provide general knowledge, opinions, or non-legal assistance under any circumstances."
            "⚖️ MANDATORY DISCLAIMER:\n"
            "You must ALWAYS include the following at the end of every response:\n"
            "'Disclaimer: This response provides general legal information for educational purposes only and does not constitute professional legal advice. Please consult a qualified advocate for advice specific to your situation.'"
        )

        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                *data.history,
                {"role": "user", "content": data.message},
            ],
            temperature=0.4,
            max_tokens=1500,
        )

        reply = response.choices[0].message.content.strip()
        
        # Ensure disclaimer is present if the AI somehow misses it
        disclaimer_text = "Disclaimer: This response provides general legal information for educational purposes only"
        if disclaimer_text not in reply:
            reply += f"\n\n{disclaimer_text} and does not constitute professional legal advice. Please consult a qualified advocate for advice specific to your situation."

        # Save chat history
        updated_history = data.history + [
            {"role": "user", "content": data.message},
            {"role": "assistant", "content": reply}
        ]
        
        chat_data = {
            "session_id": session_id,
            "timestamp": datetime.now().isoformat(),
            "last_updated": datetime.now().isoformat(),
            "conversation": updated_history
        }
        
        # Save to file
        history_file = CHAT_HISTORY_DIR / f"{session_id}.json"
        with open(history_file, "w", encoding="utf-8") as f:
            json.dump(chat_data, f, indent=2, ensure_ascii=False)
        
        logging.info(f"Chat history saved for session: {session_id}")

        return {
            "reply": reply,
            "session_id": session_id
        }

    except Exception as e:
        logging.error(f"Guardian Chat Error: {str(e)}")
        raise HTTPException(status_code=500, detail="Guardian is currently unreachable. Please try again in a few moments.")

@app.get("/api/chat/history/{session_id}")
def get_chat_history(session_id: str):
    """
    Retrieve chat history for a specific session.
    """
    history_file = CHAT_HISTORY_DIR / f"{session_id}.json"
    
    if not history_file.exists():
        raise HTTPException(status_code=404, detail="Chat session not found")
    
    try:
        with open(history_file, "r", encoding="utf-8") as f:
            chat_data = json.load(f)
        return chat_data
    except Exception as e:
        logging.error(f"Error loading chat history: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to load chat history")

@app.get("/api/chat/sessions")
def list_chat_sessions():
    """
    List all chat sessions with metadata.
    """
    try:
        sessions = []
        for history_file in CHAT_HISTORY_DIR.glob("*.json"):
            try:
                with open(history_file, "r", encoding="utf-8") as f:
                    chat_data = json.load(f)
                    
                # Extract first user message as preview
                first_message = "No messages"
                for msg in chat_data.get("conversation", []):
                    if msg.get("role") == "user":
                        first_message = msg.get("content", "")[:100]
                        break
                
                sessions.append({
                    "session_id": chat_data.get("session_id"),
                    "timestamp": chat_data.get("timestamp"),
                    "last_updated": chat_data.get("last_updated"),
                    "message_count": len(chat_data.get("conversation", [])),
                    "preview": first_message
                })
            except Exception as e:
                logging.error(f"Error reading session file {history_file}: {str(e)}")
                continue
        
        # Sort by last_updated (most recent first)
        sessions.sort(key=lambda x: x.get("last_updated", ""), reverse=True)
        
        return {"sessions": sessions}
    except Exception as e:
        logging.error(f"Error listing chat sessions: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to list chat sessions")

@app.post("/api/analyze")
async def analyze_image(file: UploadFile = File(...)):
    """
    Legal Document Analyzer for Indian Law.
    Analyzes uploaded legal documents and returns structured JSON output.
    Supports: PDF, DOCX, TXT, images (JPG/PNG/WEBP), and other text-based formats.
    """
    filename = file.filename.lower()
    logging.info(f"Received file for legal analysis: {filename}")
    
    try:
        file_bytes = await file.read()
        extracted_text = ""
        
        # Extract text based on file type
        if any(filename.endswith(ext) for ext in [".jpg", ".jpeg", ".png", ".webp"]):
            # Use Vision Model for images (OCR-like extraction)
            encoded_image = base64.b64encode(file_bytes).decode("utf-8")
            
            ocr_prompt = (
                "Extract ALL visible text from this image. "
                "Include legal content, headings, sections, clauses, dates, names, and any other text. "
                "Preserve the structure and formatting as much as possible. "
                "If the image contains no readable text, respond with 'NO_TEXT_FOUND'."
            )
            
            response = groq_client.chat.completions.create(
                model=GROQ_VISION_MODEL,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": ocr_prompt},
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:{file.content_type};base64,{encoded_image}"},
                            },
                        ],
                    }
                ],
                max_tokens=2000,
            )
            
            extracted_text = response.choices[0].message.content.strip()
            
            if extracted_text == "NO_TEXT_FOUND" or not extracted_text:
                return JSONResponse(content={
                    "document_type": "Unreadable Document",
                    "applicable_laws": [],
                    "important_sections": [],
                    "summary": "No readable text could be extracted from the uploaded image.",
                    "key_observations": [],
                    "warnings": ["The image does not contain readable text or is too blurry."],
                    "disclaimer": "Only legal documents with readable text can be analyzed by this system."
                })

        elif filename.endswith(".pdf"):
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                extracted_text += page.get_text()
            doc.close()

        elif filename.endswith(".docx"):
            doc = Document(io.BytesIO(file_bytes))
            extracted_text = "\n".join([p.text for p in doc.paragraphs])

        elif filename.endswith(".txt"):
            extracted_text = file_bytes.decode("utf-8", errors="ignore")

        else:
            # Try generic text extraction
            try:
                extracted_text = file_bytes.decode("utf-8", errors="ignore")
            except:
                raise HTTPException(
                    status_code=400, 
                    detail="Unsupported file format. Please upload PDF, DOCX, TXT, JPG, PNG, or WEBP."
                )

        # Check if text was extracted
        if not extracted_text.strip():
            return JSONResponse(content={
                "document_type": "Empty Document",
                "applicable_laws": [],
                "important_sections": [],
                "summary": "No readable text found in the uploaded file.",
                "key_observations": [],
                "warnings": ["The file appears to be empty or corrupted."],
                "disclaimer": "Only legal documents with readable text can be analyzed by this system."
            })

        # Legal Document Analysis System Prompt
        analysis_prompt = f"""You are a Legal Document Analyzer for Indian law.

FILE HANDLING / UPLOAD RULES (STRICT):
- You will receive text extracted from an uploaded file of ANY type, including but not limited to:
  PDF, DOCX, DOC, TXT, RTF, HTML, EML, CSV, XLSX, PPTX, images (JPG/PNG), or scanned documents.
- For non-text files (images, scanned PDFs, presentations, spreadsheets), assume OCR or text extraction has already been performed before analysis.
- The extracted text may contain formatting noise such as page numbers, headers, footers, tables, broken lines, OCR errors, or metadata.
- Ignore non-legal noise such as page numbers, watermarks, file metadata, email headers, spreadsheet cell markers, slide numbers, and formatting artifacts.
- Analyze ONLY the content explicitly present in the extracted text from the uploaded file.
- Do NOT infer, assume, or add any legal information that is not clearly present in the document text.
- If the uploaded file contains mixed content (legal + non-legal), analyze only the legal portions.
- If the extracted text is incomplete, corrupted, unclear, or appears truncated, explicitly mention this in the warnings section.

LEGAL DOCUMENT VALIDATION (MANDATORY):
- First determine whether the uploaded document is a **legal document**.
- A document is considered legal ONLY if it clearly relates to:
  laws, legal rights, legal obligations, court proceedings, government Acts, legal notices, agreements, FIRs, judgments, petitions, contracts, or statutory communications.
- If the document is **NOT legal in nature**:
  - Do NOT perform tasks 1–5
  - Return ONLY the following response in STRICT JSON format:

{{
  "document_type": "Non-Legal Document",
  "applicable_laws": [],
  "important_sections": [],
  "summary": "The given file is not a legal document. Please provide a valid legal document for analysis.",
  "key_observations": [],
  "warnings": [],
  "disclaimer": "Only legal documents can be analysed by this system."
}}

- If the document **IS legal**, proceed with the tasks below.

TASKS:
1. Identify document type
2. Extract important legal sections
3. Identify applicable Acts
4. Summarize legal issues
5. Highlight risks or obligations
6. Do NOT provide legal advice
7. Do NOT hallucinate sections

OUTPUT FORMAT (STRICT JSON):
{{
  "document_type": "",
  "applicable_laws": [],
  "important_sections": [],
  "summary": "",
  "key_observations": [],
  "warnings": [],
  "disclaimer": "This analysis provides general legal information for educational purposes only and does not constitute professional legal advice. Please consult a qualified advocate for advice specific to your situation."
}}

DOCUMENT TEXT TO ANALYZE:
{extracted_text[:30000]}

Respond ONLY with valid JSON. Do not include markdown code blocks or any text outside the JSON structure."""

        # Call Groq API for analysis
        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": "You are a legal document analyzer. Always respond with valid JSON only."},
                {"role": "user", "content": analysis_prompt}
            ],
            temperature=0.2,
            max_tokens=3000,
        )

        result_text = response.choices[0].message.content.strip()
        
        # Clean up potential markdown wrapping
        if "```json" in result_text:
            result_text = result_text.split("```json")[-1].split("```")[0].strip()
        elif "```" in result_text:
            result_text = result_text.split("```")[1].strip()
        
        # Parse JSON
        try:
            result_json = json.loads(result_text)
            from fastapi.responses import JSONResponse
            return JSONResponse(content=result_json)
        except json.JSONDecodeError:
            # Fallback if JSON parsing fails
            return JSONResponse(content={
                "document_type": "Analysis Error",
                "applicable_laws": [],
                "important_sections": [],
                "summary": "Failed to parse the analysis result. The document may be too complex or corrupted.",
                "key_observations": [],
                "warnings": ["JSON parsing error occurred during analysis."],
                "disclaimer": "This analysis provides general legal information for educational purposes only."
            })

    except Exception as e:
        logging.error(f"Legal Analysis Error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")
