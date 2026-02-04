from fastapi import FastAPI, HTTPException, Query, UploadFile, File, Form, Depends
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from typing import List
import os
import base64
from pathlib import Path
from docx import Document
from pydantic import BaseModel
from groq import Groq
import uuid
from datetime import datetime
from dotenv import load_dotenv
import json
import logging
import fitz  # PyMuPDF
from PIL import Image as PILImage
import io
import uvicorn
from sqlalchemy.orm import Session

# Import existing routes and DB logic
from routes.authenticator import router as auth_router
from db.db import create_tables, get_db

# Load environment variables
load_dotenv()

# Setup logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

app = FastAPI(title="Guardian AI API", version="2.0.0")

# CORS Configuration
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Constants
GROQ_MODEL = "llama-3.3-70b-versatile"
GROQ_VISION_MODEL = "llama-3.2-11b-vision-preview"

# Path definitions (Unified)
# We assume the server runs from the 'Backend' directory
BASE_DIR = Path(__file__).resolve().parent
CHATBOT_DIR = BASE_DIR.parent / "Chatbot" / "backend"
TEMPLATES_DIR = CHATBOT_DIR / "templates"
CHAT_HISTORY_DIR = CHATBOT_DIR / "chat_history"

# Ensure directories exist
CHAT_HISTORY_DIR.mkdir(parents=True, exist_ok=True)

# Clients
groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# Create tables on startup
try:
    create_tables()
except Exception as e:
    logging.warning(f"Could not create tables on startup: {e}")

# Include Authentication Router
app.include_router(auth_router, prefix="/auth", tags=["Authentication"])

# --- Models from Chatbot ---
class AIStartRequest(BaseModel):
    category: str
    subtype: str | None = None
    user_input: str

class ChatRequest(BaseModel):
    message: str
    history: List[dict] = []
    session_id: str | None = None

class AINextRequest(BaseModel):
    category: str
    filename: str
    messages: List[dict]

class AICompleteRequest(BaseModel):
    category: str
    filename: str
    messages: List[dict]

# --- Helper Functions ---
def load_metadata(category: str, subtype: str | None = None):
    if subtype:
        target_path = TEMPLATES_DIR / category / subtype / "metadata.json"
    else:
        target_path = TEMPLATES_DIR / category / "metadata.json"
    
    if not target_path.exists():
        logging.error(f"Metadata file not found at expected path: {target_path.absolute()}")
        raise HTTPException(status_code=404, detail="metadata.json not found")
    
    try:
        with open(target_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="metadata.json is not a valid JSON file")

# --- Chatbot API Endpoints ---

@app.get("/api/templates/{category}")
def list_templates(category: str):
    category_path = TEMPLATES_DIR / category
    if not category_path.exists() or not category_path.is_dir():
        raise HTTPException(status_code=404, detail="Category not found")
    files = [f.name for f in category_path.glob("*.docx")]
    return {"templates": files}

@app.get("/api/template")
def get_template(category: str = Query(...), name: str = Query(...)):
    file_path = TEMPLATES_DIR / category / name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    with open(file_path, "rb") as f:
        encoded = base64.b64encode(f.read()).decode("utf-8")
    return {"base64": encoded}

@app.post("/api/ai/start")
def start_ai_flow(data: AIStartRequest):
    try:
        templates = load_metadata(data.category, data.subtype)
        summaries_str = "\n\n".join([f"Title: {t['title']}\nSummary: {t['summary']}\nFilename: {t['filename']}" for t in templates])

        selection_prompt = (
            "You are an expert legal assistant helping users find the most suitable document template.\n"
            f"A user described their issue as:\n\n{data.user_input.strip()}\n\n"
            "Here are the available legal document templates:\n\n"
            f"{summaries_str}\n\n"
            "Your task: Select the single most suitable template. Respond ONLY with the filename. If unsure, pick the closest reasonable match."
        )

        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": selection_prompt}],
            temperature=0.3,
            max_tokens=60,
        )

        selected_filename = response.choices[0].message.content.strip()
        matched_template = next((t for t in templates if t["filename"].strip().lower() == selected_filename.lower()), None)
        if not matched_template:
            raise HTTPException(status_code=404, detail="Template match not found")

        doc_filename = f"{selected_filename}.docx"
        doc_path = (TEMPLATES_DIR / data.category / data.subtype / doc_filename) if data.subtype else (TEMPLATES_DIR / data.category / doc_filename)

        if not doc_path.exists():
            raise HTTPException(status_code=404, detail="Template file not found")

        doc = Document(doc_path)
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        question_prompt = (
            "You are a professional legal assistant helping a user complete a legal document.\n"
            "Ask questions one at a time to gather the exact information needed to fill in these blanks."
        )

        q_response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "system", "content": question_prompt}, {"role": "user", "content": full_text}],
            temperature=0.3,
            max_tokens=150,
        )

        return {
            "question": q_response.choices[0].message.content.strip(),
            "filename": f"{data.subtype}/{selected_filename}" if data.subtype else selected_filename
        }
    except Exception as e:
        logging.exception("Error in /api/ai/start")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/ai/next")
def ai_next_question(data: AINextRequest):
    try:
        subtype_path = data.filename.split("/", 1)
        if len(subtype_path) == 2:
            subfolder, filename = subtype_path
            file_path = TEMPLATES_DIR / data.category / subfolder / f"{filename}.docx"
        else:
            file_path = TEMPLATES_DIR / data.category / f"{data.filename}.docx"

        doc = Document(file_path)
        template_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        system_prompt = "You are a legal assistant. Ask ONE specific question at a time to fill placeholders. Use __COMPLETE__ if finished."
        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Template:\n{template_text}"},
                *data.messages
            ],
            temperature=0.3,
            max_tokens=300
        )
        return {"reply": response.choices[0].message.content.strip()}
    except Exception as e:
        logging.exception("Error in /api/ai/next")
        return {"reply": "Error connecting to AI service."}

@app.post("/api/chat")
def legal_chat(data: ChatRequest):
    try:
        session_id = data.session_id or str(uuid.uuid4())
        system_prompt = (
            "You are Guardian, a Legal Information Assistant for Indian law. "
            "Provide clear, accurate legal information. Do not give professional legal advice. "
            "Always include a disclaimer at the end."
        )

        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "system", "content": system_prompt}, *data.history, {"role": "user", "content": data.message}],
            temperature=0.4,
            max_tokens=1500,
        )

        reply = response.choices[0].message.content.strip()
        disclaimer = "Disclaimer: This response provides general legal information for educational purposes only."
        if disclaimer not in reply:
            reply += f"\n\n{disclaimer}"

        # Save History
        history_file = CHAT_HISTORY_DIR / f"{session_id}.json"
        with open(history_file, "w", encoding="utf-8") as f:
            json.dump({
                "session_id": session_id,
                "timestamp": datetime.now().isoformat(),
                "conversation": data.history + [{"role": "user", "content": data.message}, {"role": "assistant", "content": reply}]
            }, f, indent=2)

        return {"reply": reply, "session_id": session_id}
    except Exception as e:
        logging.exception("Error in /api/chat")
        raise HTTPException(status_code=500, detail="Chat service error")

@app.post("/api/analyze")
async def analyze_document(file: UploadFile = File(...)):
    try:
        file_bytes = await file.read()
        extracted_text = ""
        filename = file.filename.lower()

        if any(filename.endswith(ext) for ext in [".jpg", ".jpeg", ".png", ".webp"]):
            encoded_image = base64.b64encode(file_bytes).decode("utf-8")
            response = groq_client.chat.completions.create(
                model=GROQ_VISION_MODEL,
                messages=[{"role": "user", "content": [{"type": "text", "text": "Extract all text from this legal document image."}, {"type": "image_url", "image_url": {"url": f"data:{file.content_type};base64,{encoded_image}"}}]}],
                max_tokens=2000,
            )
            extracted_text = response.choices[0].message.content.strip()
        elif filename.endswith(".pdf"):
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            extracted_text = "".join([page.get_text() for page in doc])
            doc.close()
        elif filename.endswith(".docx"):
            doc = Document(io.BytesIO(file_bytes))
            extracted_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            extracted_text = file_bytes.decode("utf-8", errors="ignore")

        analysis_prompt = f"Analyze this legal document text and return ONLY JSON:\n{extracted_text[:10000]}"
        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "system", "content": "You are a legal analyzer. Output JSON only."}, {"role": "user", "content": analysis_prompt}],
            temperature=0.2,
            max_tokens=3000,
        )
        
        result_text = response.choices[0].message.content.strip()
        if "```json" in result_text:
            result_text = result_text.split("```json")[-1].split("```")[0].strip()
        
        return JSONResponse(content=json.loads(result_text))
    except Exception as e:
        logging.exception("Error in /api/analyze")
        raise HTTPException(status_code=500, detail=str(e))

# Static File Serving
if (CHATBOT_DIR.parent / "Frontend").exists():
    app.mount("/chatbot", StaticFiles(directory=str(CHATBOT_DIR.parent / "Frontend")), name="chatbot")

if (BASE_DIR.parent / "Frontend").exists():
    app.mount("/frontend", StaticFiles(directory=str(BASE_DIR.parent / "Frontend")), name="frontend")


if __name__ == "__main__":
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)

